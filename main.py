#!/usr/bin/env python3
"""DocConvert - Universal document converter: PDF, Word, Image inter-conversion."""

import argparse
import os
import sys
import tempfile
from pathlib import Path

try:
    import pymupdf
    sys.modules['fitz'] = pymupdf
    from pdf2docx import Converter
except ImportError:
    try:
        from pdf2docx import Converter
    except ImportError:
        print("Error: pdf2docx not installed.")
        print("Run: pip install pdf2docx")
        sys.exit(1)

try:
    from PIL import Image
except ImportError:
    print("Error: Pillow not installed.")
    print("Run: pip install Pillow")
    sys.exit(1)

try:
    import docx
except ImportError:
    print("Error: python-docx not installed.")
    print("Run: pip install python-docx")
    sys.exit(1)


SUPPORTED_INPUT = ".pdf"
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".webp"}


def convert_single(pdf_path, docx_path=None, start=0, end=None):
    """Convert a single PDF file to Word document."""
    pdf_path = Path(pdf_path)
    if not pdf_path.exists():
        print(f"Error: File not found: {pdf_path}")
        return False

    if pdf_path.suffix.lower() != SUPPORTED_INPUT:
        print(f"Error: Not a PDF file: {pdf_path}")
        return False

    if docx_path is None:
        docx_path = pdf_path.with_suffix(".docx")
    else:
        docx_path = Path(docx_path)

    print(f"Converting: {pdf_path.name} -> {docx_path.name}")

    try:
        cv = Converter(str(pdf_path))
        cv.convert(str(docx_path), start=start, end=end)
        cv.close()
        print(f"  Done! Saved to: {docx_path}")
        return True
    except Exception as e:
        print(f"  Error converting {pdf_path.name}: {e}")
        return False


def convert_batch(folder_path, output_folder=None, start=0, end=None):
    """Batch convert all PDF files in a folder."""
    folder = Path(folder_path)
    if not folder.exists():
        print(f"Error: Folder not found: {folder_path}")
        return

    pdf_files = sorted(f for f in folder.iterdir() if f.suffix.lower() == SUPPORTED_INPUT)

    if not pdf_files:
        print("No PDF files found in the folder.")
        return

    print(f"Found {len(pdf_files)} PDF file(s)\n")

    if output_folder:
        out_dir = Path(output_folder)
        out_dir.mkdir(parents=True, exist_ok=True)
    else:
        out_dir = folder

    success = 0
    failed = 0

    for i, pdf in enumerate(pdf_files, 1):
        print(f"[{i}/{len(pdf_files)}]", end=" ")
        docx_path = out_dir / pdf.with_suffix(".docx").name
        ok = convert_single(pdf, docx_path, start=start, end=end)
        if ok:
            success += 1
        else:
            failed += 1

    print(f"\nResults: {success} succeeded, {failed} failed, {len(pdf_files)} total.")


def list_pdf_info(pdf_path):
    """Show basic info about a PDF file."""
    pdf_path = Path(pdf_path)
    if not pdf_path.exists():
        print(f"Error: File not found: {pdf_path}")
        return

    try:
        cv = Converter(str(pdf_path))
        print(f"File: {pdf_path.name}")
        print(f"Pages: {cv.page_count}")
        cv.close()
    except Exception as e:
        print(f"Error reading file: {e}")


def pdf_to_images(pdf_path, output_dir=None, dpi=150, fmt="png"):
    """Convert PDF pages to images using PyMuPDF."""
    pdf_path = Path(pdf_path)
    if not pdf_path.exists():
        print(f"Error: File not found: {pdf_path}")
        return

    if output_dir:
        out_dir = Path(output_dir)
    else:
        out_dir = pdf_path.parent / f"{pdf_path.stem}_images"
    out_dir.mkdir(parents=True, exist_ok=True)

    try:
        doc = pymupdf.open(str(pdf_path))
        print(f"PDF: {pdf_path.name} ({doc.page_count} pages)")
        print(f"Output: {out_dir} (format: {fmt}, dpi: {dpi})\n")

        zoom = dpi / 72
        mat = pymupdf.Matrix(zoom, zoom)

        total = doc.page_count

        for i in range(total):
            page = doc[i]
            pix = page.get_pixmap(matrix=mat)
            img_name = f"{pdf_path.stem}_p{i+1:03d}.{fmt}"
            img_path = out_dir / img_name
            pix.save(str(img_path))
            print(f"  [{i+1}/{total}] {img_name}")

        doc.close()
        print(f"\nDone! {total} images saved to {out_dir}")
    except Exception as e:
        print(f"Error: {e}")


def images_to_pdf(input_path, output_path=None):
    """Combine multiple images into a single PDF using Pillow."""
    input_path = Path(input_path)

    if input_path.is_dir():
        imgs = sorted([f for f in input_path.iterdir() if f.suffix.lower() in IMAGE_EXTENSIONS])
    elif input_path.suffix.lower() in IMAGE_EXTENSIONS:
        imgs = [input_path]
    else:
        print(f"Error: Not an image or directory: {input_path}")
        return

    if not imgs:
        print("No images found.")
        return

    if output_path:
        out = Path(output_path)
    else:
        out = input_path.parent / "combined.pdf" if input_path.is_dir() else input_path.with_suffix(".pdf")

    print(f"Images: {len(imgs)}")
    print(f"Output: {out}\n")

    try:
        pil_images = []
        for i, img_path in enumerate(imgs, 1):
            img = Image.open(str(img_path))
            if img.mode == "RGBA":
                img = img.convert("RGB")
            pil_images.append(img)
            print(f"  [{i}/{len(imgs)}] {img_path.name} ({img.size[0]}x{img.size[1]})")

        pil_images[0].save(str(out), "PDF", save_all=True, append_images=pil_images[1:])
        print(f"\nDone! {len(imgs)} images combined into {out}")
    except Exception as e:
        print(f"Error: {e}")


def word_to_pdf(word_path, output_path=None):
    """Convert Word document to PDF using Microsoft Word COM automation."""
    word_path = Path(word_path)
    if not word_path.exists():
        print(f"Error: File not found: {word_path}")
        return

    if word_path.suffix.lower() not in (".doc", ".docx"):
        print(f"Error: Not a Word file: {word_path}")
        return

    if output_path:
        out = Path(output_path)
    else:
        out = word_path.with_suffix(".pdf")

    print(f"Word: {word_path.name}")
    print(f"Output: {out}\n")

    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(word_path.resolve()))
        doc.SaveAs(str(out.resolve()), FileFormat=17)
        doc.Close()
        word.Quit()
        print(f"Done! {out}")
    except ImportError:
        print("Error: pywin32 not installed.")
        print("Run: pip install pywin32")
    except Exception as e:
        print(f"Error: {e}")
        print("Note: Microsoft Word must be installed for Word → PDF conversion.")


def images_to_word(input_path, output_path=None):
    """Insert multiple images into a Word document using python-docx."""
    input_path = Path(input_path)

    if input_path.is_dir():
        imgs = sorted([f for f in input_path.iterdir() if f.suffix.lower() in IMAGE_EXTENSIONS])
    elif input_path.suffix.lower() in IMAGE_EXTENSIONS:
        imgs = [input_path]
    else:
        print(f"Error: Not an image or directory: {input_path}")
        return

    if not imgs:
        print("No images found.")
        return

    if output_path:
        out = Path(output_path)
    else:
        out = input_path.parent / "images.docx" if input_path.is_dir() else input_path.with_suffix(".docx")

    print(f"Images: {len(imgs)}")
    print(f"Output: {out}\n")

    try:
        doc = docx.Document()
        for i, img_path in enumerate(imgs, 1):
            doc.add_picture(str(img_path), width=docx.shared.Inches(6))
            doc.add_paragraph(f"Image {i}: {img_path.name}")
            print(f"  [{i}/{len(imgs)}] {img_path.name}")
        doc.save(str(out))
        print(f"\nDone! {len(imgs)} images inserted into {out}")
    except Exception as e:
        print(f"Error: {e}")


def word_to_images(word_path, output_dir=None, dpi=150, fmt="png"):
    """Convert Word to images via PDF (Word → PDF → Images)."""
    word_path = Path(word_path)
    if not word_path.exists():
        print(f"Error: File not found: {word_path}")
        return

    if output_dir:
        out_dir = Path(output_dir)
    else:
        out_dir = word_path.parent / f"{word_path.stem}_images"
    out_dir.mkdir(parents=True, exist_ok=True)

    print("Step 1: Word → PDF")
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False, dir=str(out_dir)) as tmp:
        tmp_pdf = Path(tmp.name)

    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(word_path.resolve()))
        doc.SaveAs(str(tmp_pdf.resolve()), FileFormat=17)
        doc.Close()
        word.Quit()
    except Exception as e:
        print(f"  Word → PDF failed: {e}")
        print("  Note: Microsoft Word must be installed.")
        if tmp_pdf.exists():
            tmp_pdf.unlink()
        return

    print("Step 2: PDF → Images")
    pdf_to_images(tmp_pdf, out_dir, dpi=dpi, fmt=fmt)
    tmp_pdf.unlink()


def pdf_to_text(pdf_path, output_path=None):
    """Extract text from a PDF file using PyMuPDF."""
    pdf_path = Path(pdf_path)
    if not pdf_path.exists():
        print(f"Error: File not found: {pdf_path}")
        return

    if output_path:
        out = Path(output_path)
    else:
        out = pdf_path.with_suffix(".txt")

    try:
        doc = pymupdf.open(str(pdf_path))
        total = doc.page_count
        text_parts = []
        for i in range(total):
            page = doc[i]
            text_parts.append(f"--- Page {i+1} ---\n")
            text_parts.append(page.get_text())
        doc.close()

        out.write_text("".join(text_parts), encoding="utf-8")
        print(f"PDF: {pdf_path.name} ({total} pages)")
        print(f"Output: {out}")
        print(f"Done! Text extracted to {out}")
    except Exception as e:
        print(f"Error: {e}")


def convert_image_format(input_path, output_path=None, fmt="png"):
    """Convert image to a different format using Pillow."""
    input_path = Path(input_path)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        return

    if input_path.is_dir():
        imgs = sorted([f for f in input_path.iterdir() if f.suffix.lower() in IMAGE_EXTENSIONS])
    elif input_path.suffix.lower() in IMAGE_EXTENSIONS:
        imgs = [input_path]
    else:
        print(f"Error: Not an image or directory: {input_path}")
        return

    if not imgs:
        print("No images found.")
        return

    out_dir = Path(output_path) if output_path else input_path.parent if input_path.is_file() else input_path
    if not out_dir.exists():
        out_dir.mkdir(parents=True, exist_ok=True)

    print(f"Converting {len(imgs)} image(s) to {fmt.upper()}\n")
    success = 0
    for i, img_path in enumerate(imgs, 1):
        try:
            img = Image.open(str(img_path))
            out_name = f"{img_path.stem}.{fmt}"
            out_path = out_dir / out_name
            if fmt == "jpg" and img.mode == "RGBA":
                img = img.convert("RGB")
            img.save(str(out_path), fmt.upper() if fmt != "jpg" else "JPEG")
            print(f"  [{i}/{len(imgs)}] {img_path.name} -> {out_name}")
            success += 1
        except Exception as e:
            print(f"  [{i}/{len(imgs)}] {img_path.name} -> Error: {e}")

    print(f"\nDone! {success}/{len(imgs)} converted.")


def merge_pdfs(pdf_list, output_path):
    """Merge multiple PDF files into one using PyMuPDF."""
    pdfs = [Path(p) for p in pdf_list if Path(p).exists()]
    if len(pdfs) < 2:
        print("Error: Need at least 2 PDF files to merge.")
        return

    out = Path(output_path) if output_path else Path("merged.pdf")

    try:
        result = pymupdf.open()
        for i, pdf in enumerate(pdfs, 1):
            src = pymupdf.open(str(pdf))
            result.insert_pdf(src)
            print(f"  [{i}/{len(pdfs)}] {pdf.name} ({src.page_count} pages)")
            src.close()

        result.save(str(out))
        total = result.page_count
        result.close()
        print(f"\nDone! {len(pdfs)} PDFs merged into {out} ({total} pages)")
    except Exception as e:
        print(f"Error: {e}")


def split_pdf(pdf_path, output_dir=None, pages_per_file=1):
    """Split a PDF into multiple files using PyMuPDF."""
    pdf_path = Path(pdf_path)
    if not pdf_path.exists():
        print(f"Error: File not found: {pdf_path}")
        return

    if output_dir:
        out_dir = Path(output_dir)
    else:
        out_dir = pdf_path.parent / f"{pdf_path.stem}_split"
    out_dir.mkdir(parents=True, exist_ok=True)

    try:
        doc = pymupdf.open(str(pdf_path))
        total = doc.page_count
        print(f"PDF: {pdf_path.name} ({total} pages)")
        print(f"Output: {out_dir} ({pages_per_file} pages/file)\n")

        part = 0
        start = 0
        while start < total:
            end = min(start + pages_per_file, total)
            out_path = out_dir / f"{pdf_path.stem}_part{part+1:03d}.pdf"
            new_doc = pymupdf.open()
            new_doc.insert_pdf(doc, from_page=start, to_page=end - 1)
            new_doc.save(str(out_path))
            new_doc.close()
            print(f"  Part {part+1}: pages {start+1}-{end} -> {out_path.name}")
            start = end
            part += 1

        doc.close()
        print(f"\nDone! Split into {part} files.")
    except Exception as e:
        print(f"Error: {e}")


def excel_to_pdf(excel_path, output_path=None):
    """Convert Excel file to PDF using Microsoft Excel COM automation."""
    excel_path = Path(excel_path)
    if not excel_path.exists():
        print(f"Error: File not found: {excel_path}")
        return

    if excel_path.suffix.lower() not in (".xls", ".xlsx"):
        print(f"Error: Not an Excel file: {excel_path}")
        return

    if output_path:
        out = Path(output_path)
    else:
        out = excel_path.with_suffix(".pdf")

    print(f"Excel: {excel_path.name}")
    print(f"Output: {out}\n")

    try:
        import win32com.client
        excel = win32com.client.Dispatch("Excel.Application")
        excel.Visible = False
        wb = excel.Workbooks.Open(str(excel_path.resolve()))
        wb.ExportAsFixedFormat(0, str(out.resolve()))
        wb.Close(False)
        excel.Quit()
        print(f"Done! {out}")
    except Exception as e:
        print(f"Error: {e}")
        print("Note: Microsoft Excel must be installed.")


def ppt_to_pdf(ppt_path, output_path=None):
    """Convert PowerPoint file to PDF using Microsoft PowerPoint COM automation."""
    ppt_path = Path(ppt_path)
    if not ppt_path.exists():
        print(f"Error: File not found: {ppt_path}")
        return

    if ppt_path.suffix.lower() not in (".ppt", ".pptx"):
        print(f"Error: Not a PowerPoint file: {ppt_path}")
        return

    if output_path:
        out = Path(output_path)
    else:
        out = ppt_path.with_suffix(".pdf")

    print(f"PowerPoint: {ppt_path.name}")
    print(f"Output: {out}\n")

    try:
        import win32com.client
        ppt = win32com.client.Dispatch("PowerPoint.Application")
        ppt.Visible = True
        pres = ppt.Presentations.Open(str(ppt_path.resolve()), WithWindow=False)
        pres.SaveAs(str(out.resolve()), 32)
        pres.Close()
        ppt.Quit()
        print(f"Done! {out}")
    except Exception as e:
        print(f"Error: {e}")
        print("Note: Microsoft PowerPoint must be installed.")


def text_to_pdf(text_path, output_path=None):
    """Convert a text file to PDF using fpdf2."""
    text_path = Path(text_path)
    if not text_path.exists():
        print(f"Error: File not found: {text_path}")
        return

    if output_path:
        out = Path(output_path)
    else:
        out = text_path.with_suffix(".pdf")

    print(f"Text: {text_path.name}")
    print(f"Output: {out}\n")

    try:
        from fpdf import FPDF, XPos, YPos

        text = text_path.read_text(encoding="utf-8")
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("helvetica", size=12)

        for line in text.split("\n"):
            try:
                pdf.cell(0, 8, text=line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            except Exception:
                pdf.cell(0, 8, text=line.encode("ascii", "replace").decode("ascii"),
                         new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        pdf.output(str(out))
        print(f"Done! {out}")
    except Exception as e:
        print(f"Error: {e}")


def main():
    parser = argparse.ArgumentParser(
        description="DocConvert - Universal document converter (PDF/Word/Image/Excel/PPT inter-conversion).",
        epilog="Examples:\n"
               "  PDF → Word:   python main.py convert document.pdf\n"
               "  Word → PDF:   python main.py word2pdf document.docx\n"
               "  PDF → Image:  python main.py pdf2img document.pdf\n"
               "  Image → PDF:  python main.py img2pdf photo.jpg\n"
               "  Image → Word: python main.py img2word photo.jpg\n"
               "  Word → Image: python main.py word2img document.docx\n"
               "  PDF → Text:   python main.py pdf2text document.pdf\n"
               "  Image format: python main.py imgconv photo.png --format jpg\n"
               "  Merge PDFs:   python main.py pdfmerge a.pdf b.pdf c.pdf\n"
               "  Split PDF:    python main.py pdfsplit document.pdf\n"
               "  Excel → PDF:  python main.py excel2pdf report.xlsx\n"
               "  PPT → PDF:    python main.py ppt2pdf slides.pptx\n"
               "  Text → PDF:   python main.py txt2pdf notes.txt\n"
               "  Batch:        python main.py batch \"C:\\Documents\"\n"
               "  PDF info:     python main.py info document.pdf",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )

    subparsers = parser.add_subparsers(dest="command", help="Available commands")

    conv_parser = subparsers.add_parser("convert", help="PDF → Word")
    conv_parser.add_argument("pdf", help="Path to the PDF file")
    conv_parser.add_argument("-o", "--output", help="Output Word file path")
    conv_parser.add_argument("--start", type=int, default=0, help="Start page (0-based)")
    conv_parser.add_argument("--end", type=int, default=None, help="End page (exclusive)")

    batch_parser = subparsers.add_parser("batch", help="Batch convert all PDFs in a folder to Word")
    batch_parser.add_argument("folder", help="Folder containing PDF files")
    batch_parser.add_argument("-o", "--output", help="Output folder")
    batch_parser.add_argument("--start", type=int, default=0, help="Start page (0-based)")
    batch_parser.add_argument("--end", type=int, default=None, help="End page (exclusive)")

    info_parser = subparsers.add_parser("info", help="Show PDF file info (page count)")
    info_parser.add_argument("pdf", help="Path to the PDF file")

    p2i_parser = subparsers.add_parser("pdf2img", help="PDF → Images (PNG/JPG)")
    p2i_parser.add_argument("pdf", help="Path to the PDF file")
    p2i_parser.add_argument("-o", "--output", help="Output directory")
    p2i_parser.add_argument("--dpi", type=int, default=150, help="Resolution DPI (default: 150)")
    p2i_parser.add_argument("--format", choices=["png", "jpg"], default="png", help="Image format (default: png)")

    i2p_parser = subparsers.add_parser("img2pdf", help="Images → PDF (single or folder)")
    i2p_parser.add_argument("input", help="Image file or folder of images")
    i2p_parser.add_argument("-o", "--output", help="Output PDF path")

    w2p_parser = subparsers.add_parser("word2pdf", help="Word → PDF (requires Microsoft Word)")
    w2p_parser.add_argument("word", help="Path to the Word file (.doc/.docx)")
    w2p_parser.add_argument("-o", "--output", help="Output PDF path")

    i2w_parser = subparsers.add_parser("img2word", help="Images → Word (single or folder)")
    i2w_parser.add_argument("input", help="Image file or folder of images")
    i2w_parser.add_argument("-o", "--output", help="Output Word path")

    w2i_parser = subparsers.add_parser("word2img", help="Word → Images (requires Microsoft Word)")
    w2i_parser.add_argument("word", help="Path to the Word file (.doc/.docx)")
    w2i_parser.add_argument("-o", "--output", help="Output directory")
    w2i_parser.add_argument("--dpi", type=int, default=150, help="Resolution DPI (default: 150)")
    w2i_parser.add_argument("--format", choices=["png", "jpg"], default="png", help="Image format (default: png)")

    p2t_parser = subparsers.add_parser("pdf2text", help="PDF → Text (extract text)")
    p2t_parser.add_argument("pdf", help="Path to the PDF file")
    p2t_parser.add_argument("-o", "--output", help="Output text file path")

    imgc_parser = subparsers.add_parser("imgconv", help="Image format conversion (JPG/PNG/WebP/BMP)")
    imgc_parser.add_argument("input", help="Image file or folder of images")
    imgc_parser.add_argument("-o", "--output", help="Output directory")
    imgc_parser.add_argument("--format", choices=["png", "jpg", "webp", "bmp"], default="png", help="Target format (default: png)")

    pm_parser = subparsers.add_parser("pdfmerge", help="Merge multiple PDFs into one")
    pm_parser.add_argument("pdfs", nargs="+", help="PDF files to merge (at least 2)")
    pm_parser.add_argument("-o", "--output", help="Output PDF path")

    ps_parser = subparsers.add_parser("pdfsplit", help="Split a PDF into multiple files")
    ps_parser.add_argument("pdf", help="Path to the PDF file")
    ps_parser.add_argument("-o", "--output", help="Output directory")
    ps_parser.add_argument("--pages", type=int, default=1, help="Pages per file (default: 1)")

    e2p_parser = subparsers.add_parser("excel2pdf", help="Excel → PDF (requires Microsoft Excel)")
    e2p_parser.add_argument("excel", help="Path to the Excel file (.xls/.xlsx)")
    e2p_parser.add_argument("-o", "--output", help="Output PDF path")

    pp2p_parser = subparsers.add_parser("ppt2pdf", help="PPT → PDF (requires Microsoft PowerPoint)")
    pp2p_parser.add_argument("ppt", help="Path to the PowerPoint file (.ppt/.pptx)")
    pp2p_parser.add_argument("-o", "--output", help="Output PDF path")

    t2p_parser = subparsers.add_parser("txt2pdf", help="Text → PDF")
    t2p_parser.add_argument("text", help="Path to the text file (.txt)")
    t2p_parser.add_argument("-o", "--output", help="Output PDF path")

    args = parser.parse_args()

    if args.command == "convert":
        convert_single(args.pdf, args.output, start=args.start, end=args.end)
    elif args.command == "batch":
        convert_batch(args.folder, args.output, start=args.start, end=args.end)
    elif args.command == "info":
        list_pdf_info(args.pdf)
    elif args.command == "pdf2img":
        pdf_to_images(args.pdf, args.output, dpi=args.dpi, fmt=args.format)
    elif args.command == "img2pdf":
        images_to_pdf(args.input, args.output)
    elif args.command == "word2pdf":
        word_to_pdf(args.word, args.output)
    elif args.command == "img2word":
        images_to_word(args.input, args.output)
    elif args.command == "word2img":
        word_to_images(args.word, args.output, dpi=args.dpi, fmt=args.format)
    elif args.command == "pdf2text":
        pdf_to_text(args.pdf, args.output)
    elif args.command == "imgconv":
        convert_image_format(args.input, args.output, fmt=args.format)
    elif args.command == "pdfmerge":
        merge_pdfs(args.pdfs, args.output)
    elif args.command == "pdfsplit":
        split_pdf(args.pdf, args.output, pages_per_file=args.pages)
    elif args.command == "excel2pdf":
        excel_to_pdf(args.excel, args.output)
    elif args.command == "ppt2pdf":
        ppt_to_pdf(args.ppt, args.output)
    elif args.command == "txt2pdf":
        text_to_pdf(args.text, args.output)
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
